"""
Convert Think Ventures markdown docs to formatted Word documents.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DOCS_DIR = r"C:\Users\Chris\Desktop\WEBSITES\THINK! VENTURES\docs"

FILES = [
    ("bylaws.md",                       "Think Ventures - Bylaws.docx"),
    ("organizational-meeting-minutes.md","Think Ventures - Organizational Meeting Minutes.docx"),
    ("nonprofit-business-plan.md",       "Think Ventures - Nonprofit Business Plan.docx"),
]

def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_table_from_md(doc, lines):
    """Parse markdown table lines and add a Word table."""
    rows = [l for l in lines if l.startswith("|") and not re.match(r"\|[\s\-|]+\|", l)]
    if not rows:
        return
    cols_data = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        cols_data.append(cells)
    num_cols = max(len(r) for r in cols_data)
    table = doc.add_table(rows=len(cols_data), cols=num_cols)
    table.style = "Table Grid"
    for r_idx, row_data in enumerate(cols_data):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= num_cols:
                break
            cell = table.cell(r_idx, c_idx)
            # Strip bold markers for display
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
            clean = re.sub(r"`(.*?)`", r"\1", clean)
            p = cell.paragraphs[0]
            run = p.add_run(clean)
            if r_idx == 0:
                run.bold = True
                run.font.size = Pt(10)
                cell._tc.get_or_add_tcPr()
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "1F3A5F")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:val"), "clear")
                cell._tc.tcPr.append(shading)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                run.font.size = Pt(10)
    doc.add_paragraph()

def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Table — collect all consecutive table lines
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            add_table_from_md(doc, table_lines)
            continue

        # Code block (skip — replace with italic note)
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1  # skip closing ```
            if code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                for cl in code_lines:
                    run = p.add_run(cl + "\n")
                    run.font.name  = "Courier New"
                    run.font.size  = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # H1
        if line.startswith("# "):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            text = line[2:].strip()
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            p = doc.add_paragraph(style="Quote")
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(11)
            i += 1
            continue

        # Bullet list
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            text = re.sub(r"`(.*?)`",       r"\1", text)
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line).strip()
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Signature line (underscores)
        if line.strip().startswith("_____"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("&nbsp;", "  "))
            run.font.name  = "Courier New"
            run.font.size  = Pt(11)
            i += 1
            continue

        # Normal paragraph — handle inline bold/italic
        p = doc.add_paragraph()
        # Split on bold markers
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            else:
                # Strip markdown links [text](url) → text
                clean = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", part)
                clean = clean.replace("&nbsp;", "  ")
                p.add_run(clean)

        i += 1

    doc.save(docx_path)
    print(f"[OK] Saved: {os.path.basename(docx_path)}")


if __name__ == "__main__":
    for md_file, docx_file in FILES:
        md_path   = os.path.join(DOCS_DIR, md_file)
        docx_path = os.path.join(DOCS_DIR, docx_file)
        print(f"Converting {md_file} ...")
        convert_md_to_docx(md_path, docx_path)
    print("\n[DONE] All documents converted! Check the docs/ folder.")
